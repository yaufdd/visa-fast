#!/usr/bin/env python3
"""
FujiTravel document generator.
Usage: python3 generate.py <group_id> <pass2_json_path> <zip_path>

Reads Pass 2 JSON, generates per-tourist .docx and .pdf files plus
two group-level documents, then zips everything to zip_path.
"""

import sys
import os
import json
import re
import copy
import shutil
import zipfile

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fillpdf import fillpdfs

# pikepdf is used ONLY to overwrite T34 (former nationalities) AFTER
# fillpdf has done the main fill. fillpdf hard-validates combo box values
# against the predefined option list and would reject "USSR" / "NO" with
# a KeyError. pikepdf has no such validation, so we let fillpdf put a
# safe placeholder there, then overwrite in a tiny post-pass.
import pikepdf

# ── MVD region code → English city name ──────────────────────────────────────
_MVD_REGION_CITY = {
    "01": "Maikop",           "02": "Ufa",               "03": "Ulan-Ude",
    "04": "Gorno-Altaysk",    "05": "Makhachkala",        "06": "Magas",
    "07": "Nalchik",          "08": "Elista",             "09": "Cherkessk",
    "10": "Petrozavodsk",     "11": "Syktyvkar",          "12": "Yoshkar-Ola",
    "13": "Saransk",          "14": "Yakutsk",            "15": "Vladikavkaz",
    "16": "Kazan",            "17": "Kyzyl",              "18": "Izhevsk",
    "19": "Abakan",           "20": "Grozny",             "21": "Cheboksary",
    "22": "Barnaul",          "23": "Krasnodar",          "24": "Krasnoyarsk",
    "25": "Vladivostok",      "26": "Stavropol",          "27": "Khabarovsk",
    "28": "Blagoveshchensk",  "29": "Arkhangelsk",        "30": "Astrakhan",
    "31": "Belgorod",         "32": "Bryansk",            "33": "Vladimir",
    "34": "Volgograd",        "35": "Vologda",            "36": "Voronezh",
    "37": "Ivanovo",          "38": "Irkutsk",            "39": "Kaliningrad",
    "40": "Kaluga",           "41": "Petropavlovsk-Kamchatsky", "42": "Kemerovo",
    "43": "Kirov",            "44": "Kostroma",           "45": "Kurgan",
    "46": "Kursk",            "47": "Saint Petersburg",   "48": "Lipetsk",
    "49": "Magadan",          "50": "Moscow",             "51": "Murmansk",
    "52": "Nizhny Novgorod",  "53": "Veliky Novgorod",    "54": "Novosibirsk",
    "55": "Omsk",             "56": "Orenburg",           "57": "Oryol",
    "58": "Penza",            "59": "Perm",               "60": "Pskov",
    "61": "Rostov-on-Don",    "62": "Ryazan",             "63": "Samara",
    "64": "Saratov",          "65": "Yuzhno-Sakhalinsk",  "66": "Yekaterinburg",
    "67": "Smolensk",         "68": "Tambov",             "69": "Tver",
    "70": "Tomsk",            "71": "Tula",               "72": "Tyumen",
    "73": "Ulyanovsk",        "74": "Chelyabinsk",        "75": "Chita",
    "76": "Yaroslavl",        "77": "Moscow",             "78": "Saint Petersburg",
    "79": "Birobidzhan",      "82": "Simferopol",         "83": "Naryan-Mar",
    "86": "Khanty-Mansiysk",  "87": "Anadyr",             "89": "Salekhard",
    "91": "Sevastopol",
}

def _mvd_to_place_of_issue(issued_by: str) -> str:
    """Convert 'MVD 54001' → 'Russia, Novosibirsk'.
    If already contains 'Russia' or city info, return as-is.
    Fallback: 'Russia'."""
    if not issued_by:
        return "Russia"
    # Already formatted nicely
    if "Russia" in issued_by or "RUSSIA" in issued_by:
        return issued_by
    # Extract first 2 digits from MVD code: "MVD 54001", "MVD54001", "54001"
    m = re.search(r'(?:MVD\s*)?(\d{2})\d+', issued_by, re.IGNORECASE)
    if m:
        city = _MVD_REGION_CITY.get(m.group(1))
        if city:
            return f"Russia, {city}"
    return "Russia"

# ── Paths ─────────────────────────────────────────────────────────────────────
TEMPLATES_DIR = os.environ.get("DOCGEN_TEMPLATES_DIR", "/Users/yaufdd/Desktop/FUJIT TRAVEL/Шаблоны")
PDF_TEMPLATE  = os.environ.get("DOCGEN_PDF_TEMPLATE", "/Users/yaufdd/Desktop/FUJIT TRAVEL/ТЕСТ_Бамба/Бамба Эрик.pdf")

TMPL_PROGRAMME   = os.path.join(TEMPLATES_DIR, "ШАБЛОН программа.docx")
# Per-org custom доверенность template — when set, replaces the bundled
# template for this generation run. Set by the Go layer when the org has
# uploaded a custom .docx via /api/templates/doverenost.
TMPL_DOVERENOST  = os.environ.get("DOCGEN_DOVERENOST_TEMPLATE") or os.path.join(TEMPLATES_DIR, "ШАБЛОН доверенность.docx")
TMPL_INNA        = os.path.join(TEMPLATES_DIR, "ШАБЛОН для Инны в ВЦ.docx")
TMPL_VC_REQUEST  = os.path.join(TEMPLATES_DIR, "ШАБЛОН заявка ВЦ.docx")


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_text(cell, text):
    """Replace all text in a table cell, preserving the first paragraph's style."""
    # Clear all paragraphs except the first
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    first = cell.paragraphs[0]
    for run in first.runs:
        run.text = ""
    if not first.runs:
        first.add_run(text)
    else:
        first.runs[0].text = text


def set_cell_multiline(cell, text):
    """
    Write text into a cell using soft line breaks (shift+enter = <w:br/>) between lines.
    Single \\n  → soft line break (shift+enter, small gap)
    Double \\n\\n → blank line (empty line break, larger gap)
    """
    # Remove all paragraphs except the first
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)

    para = cell.paragraphs[0]
    # Remove all existing runs
    for r in list(para._p.findall(qn("w:r"))):
        para._p.remove(r)

    lines = text.split("\n")
    for i, line in enumerate(lines):
        run = para.add_run(line)
        if i < len(lines) - 1:
            # Add soft line break after each line except the last
            br = OxmlElement("w:br")
            run._r.append(br)


def add_table_row(table, values):
    """Append a row to a table with 4 cells filled from values list."""
    row = table.add_row()
    for i, val in enumerate(values):
        set_cell_multiline(row.cells[i], val)
    return row


def replace_paragraph_text(doc, placeholder, replacement):
    """Replace placeholder text across all paragraphs and runs in a document."""
    for para in doc.paragraphs:
        if placeholder in para.text:
            for run in para.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, replacement)
            # Fallback: rebuild the paragraph if runs didn't cover it
            if placeholder in para.text:
                full = para.text.replace(placeholder, replacement)
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = full
                else:
                    para.add_run(full)


# ── Programme .docx ───────────────────────────────────────────────────────────

def generate_programme(data, out_path, contact_phone=""):
    doc = Document(TMPL_PROGRAMME)

    replace_paragraph_text(doc, "(ДАТА СОСТАВЛЕНИЯ ДД.ММ.ГГГГ)", data["document_date"])

    all_names = ", ".join(t["name_lat"] for t in data["tourists"])
    replace_paragraph_text(doc, "(ИМЯ ФАМИЛИЯ ЛАТИНИЦЕЙ)", all_names)

    table = doc.tables[0]
    for _ in range(len(table.rows) - 1):
        row = table.rows[-1]
        row._tr.getparent().remove(row._tr)

    first_contact_set = False
    for row in data["programme"]:
        contact = row.get("contact", "")
        # Override the first real contact (not "Same as above") with oldest male's phone
        if contact_phone and not first_contact_set and contact.lower() not in ("same as above", ""):
            contact = contact_phone
            first_contact_set = True
        elif contact_phone and not first_contact_set:
            contact = contact_phone
            first_contact_set = True
        add_table_row(table, [
            row.get("date", ""),
            row.get("activity", ""),
            contact,
            row.get("accommodation", ""),
        ])

    doc.save(out_path)


# ── Doverenost .docx ──────────────────────────────────────────────────────────

def generate_doverenost(data, dov, out_path):
    import re

    doc = Document(TMPL_DOVERENOST)

    # The template has one big paragraph with placeholders
    target = doc.paragraphs[2]  # index 2 is the main body paragraph
    full = target.text

    # Name must be surrounded by commas per Russian legal conventions. We
    # wrap the value with commas and rely on the cleanup pass below to
    # collapse any duplicate commas in case the template already had them.
    name_wrapped = f", {dov['name_ru']},"

    # Replace variable placeholders
    replacements = {
        "(ФИО ПО-РУССКИ)": name_wrapped,
        "(ДД.ММ.ГГГГ)":    dov["dob"],
        "(СЕРИЯ НОМЕР)":   f"{dov['passport_series']} {dov['passport_number']}",
        "(ОРГАН ВЫДАЧИ)":  dov["issued_by"],
        "(АДРЕС РЕГИСТРАЦИИ)": dov["reg_address"],
    }
    full_new = full
    for ph, val in replacements.items():
        full_new = full_new.replace(ph, val)

    # Replace «ДД» МЕСЯЦ ГГГГ with the issued_date value
    full_new = re.sub(r"«ДД» МЕСЯЦ ГГГГ", dov["issued_date"], full_new)

    # Collapse duplicate commas/spaces that may come from pre-existing
    # template punctuation around the name placeholder or from the
    # comma-wrapping above.
    full_new = re.sub(r",[\s,]+,", ",", full_new)
    full_new = re.sub(r",\s*,", ",", full_new)
    full_new = re.sub(r"[ \t]{2,}", " ", full_new)
    full_new = re.sub(r"\s+([,.;:])", r"\1", full_new)

    # For minors: insert "своего/своей сына/дочери, [имя ребёнка]," after "документов"
    if dov.get("is_minor") and dov.get("child_name_ru"):
        child_gender = dov.get("child_gender", "сына")  # "сына" or "дочери"
        child_name   = dov["child_name_ru"]
        possessive   = "своей" if child_gender == "дочери" else "своего"
        insert_text  = f" {possessive} {child_gender}, {child_name},"
        # Remove "своих" from template's "своих документов" (POA is only for the child)
        full_new = full_new.replace("своих документов", "документов", 1)
        # Insert after "документов"
        full_new = full_new.replace("документов", "документов" + insert_text, 1)

    for run in target.runs:
        run.text = ""
    if target.runs:
        target.runs[0].text = full_new
    else:
        target.add_run(full_new)

    doc.save(out_path)


# ── Для Инны в ВЦ .docx ──────────────────────────────────────────────────────

def generate_inna(data, out_path):
    doc = Document(TMPL_INNA)
    inna = data["inna_doc"]

    replace_paragraph_text(doc, "(ДАТА ПРИЛЁТА)", inna["submission_date"])
    replace_paragraph_text(doc, "(ФИО КЛИЕНТА 1)", inna["applicants_ru"][0])

    # Add remaining applicants after the first one
    if len(inna["applicants_ru"]) > 1:
        # Find the paragraph with the first applicant and insert after it
        target_idx = None
        for i, p in enumerate(doc.paragraphs):
            if inna["applicants_ru"][0] in p.text:
                target_idx = i
                break
        if target_idx is not None:
            for name in inna["applicants_ru"][1:]:
                new_p = OxmlElement("w:p")
                new_run = OxmlElement("w:r")
                new_t = OxmlElement("w:t")
                new_t.text = name
                new_run.append(new_t)
                new_p.append(new_run)
                doc.paragraphs[target_idx]._element.addnext(new_p)
                target_idx += 1

    doc.save(out_path)


# ── Сумма прописью (970 × N) ─────────────────────────────────────────────────

# Готовый словарь: количество человек → сумма прописью (970 × N рублей)
_AMOUNT_WORDS = {
    1:  "девятьсот семьдесят",
    2:  "одна тысяча девятьсот сорок",
    3:  "две тысячи девятьсот десять",
    4:  "три тысячи восемьсот восемьдесят",
    5:  "четыре тысячи восемьсот пятьдесят",
    6:  "пять тысяч восемьсот двадцать",
    7:  "шесть тысяч семьсот девяносто",
    8:  "семь тысяч семьсот шестьдесят",
    9:  "восемь тысяч семьсот тридцать",
    10: "девять тысяч семьсот",
    11: "десять тысяч шестьсот семьдесят",
    12: "одиннадцать тысяч шестьсот сорок",
    13: "двенадцать тысяч шестьсот десять",
    14: "тринадцать тысяч пятьсот восемьдесят",
    15: "четырнадцать тысяч пятьсот пятьдесят",
    16: "пятнадцать тысяч пятьсот двадцать",
    17: "шестнадцать тысяч четыреста девяносто",
    18: "семнадцать тысяч четыреста шестьдесят",
    19: "восемнадцать тысяч четыреста тридцать",
    20: "девятнадцать тысяч четыреста",
}

def _fmt_amount(total):
    """Format total as '4 850' (Russian thousands separator)."""
    return f"{total:,}".replace(",", "\u00a0")  # non-breaking space


# ── Заявка ВЦ .docx ───────────────────────────────────────────────────────────

def generate_vc_request(data, out_path):
    doc = Document(TMPL_VC_REQUEST)
    vc = data["vc_request"]
    count = vc["count"]
    total = count * 970
    total_str = _fmt_amount(total)
    words = _AMOUNT_WORDS.get(count, str(total))

    # Table 1 — applicants
    t1 = doc.tables[1]
    template_tr = copy.deepcopy(t1.rows[2]._tr)  # snapshot BEFORE mutating text
    set_cell_text(t1.rows[2].cells[0], vc["applicants"][0])
    # Insert each additional applicant AFTER the previously inserted row to
    # preserve the list order (not reversed).
    for i, name in enumerate(vc["applicants"][1:], start=1):
        new_row = copy.deepcopy(template_tr)
        t1.rows[2 + i - 1]._tr.addnext(new_row)
        set_cell_text(t1.rows[2 + i].cells[0], name)

    # Update count row (last row of table 1)
    count_row = t1.rows[-1]
    cell_text_new = count_row.cells[0].text.replace("_______(N)_____", str(count))
    set_cell_text(count_row.cells[0], cell_text_new)

    # Table 5 row 5.1 — quantity and total
    t5 = doc.tables[5]
    row51 = t5.rows[1]
    set_cell_text(row51.cells[3], str(count))
    set_cell_text(row51.cells[4], total_str)

    # Replace (N) and (N × 970) in table 5 cells
    for row in t5.rows:
        for cell in row.cells:
            if "(N × 970)" in cell.text:
                set_cell_text(cell, total_str)
            elif cell.text.strip() == "(N)":
                set_cell_text(cell, str(count))

    # Replace in all paragraphs: "(N × 970) (СУММА ПРОПИСЬЮ)" → "4 850 (четыре тысячи...)"
    for para in doc.paragraphs:
        if "(N × 970)" in para.text or "(СУММА ПРОПИСЬЮ)" in para.text:
            new_text = para.text
            new_text = new_text.replace("(N × 970) (СУММА ПРОПИСЬЮ)", f"{total_str} ({words})")
            new_text = new_text.replace("(N × 970)", total_str)
            new_text = new_text.replace("(СУММА ПРОПИСЬЮ)", f"({words})")
            new_text = new_text.replace("(N)", str(count))
            # Rebuild paragraph preserving first run's formatting
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = new_text
            else:
                para.add_run(new_text)

    doc.save(out_path)


# ── Анкета PDF ────────────────────────────────────────────────────────────────

def generate_anketa(tourist, anketa, dov, out_path, departure_date_str=""):
    # Split name_lat into first/last
    parts = tourist.get("name_lat", "").split()
    last_name  = parts[0] if parts else ""
    first_name = " ".join(parts[1:]) if len(parts) > 1 else ""

    # "Other names" (T16[1]) — maiden name from sheet; "NO" if empty.
    # Pass 2 transliterates the surname to Latin if present.
    maiden_name_val = tourist.get("maiden_name", "") or "NO"
    previous_visits_val = tourist.get("previous_visits", "") or "NO"

    # Compute intended stay days from THIS tourist's own flight dates.
    # Stay = departure_date - arrival_date + 1 (both endpoints inclusive).
    from datetime import datetime as _dt
    stay_days = tourist.get("intended_stay_days", 0)
    try:
        arrival_str = tourist.get("arrival_date_japan", "")
        dep_str = tourist.get("departure_date_japan", "") or departure_date_str
        if arrival_str and dep_str:
            arr = _dt.strptime(arrival_str, "%d.%m.%Y")
            dep = _dt.strptime(dep_str, "%d.%m.%Y")
            stay_days = (dep - arr).days + 1
    except Exception:
        pass  # keep AI value as fallback

    fields = {
        # Page 1
        "topmostSubform[0].Page1[0].T2[0]":  last_name,
        "topmostSubform[0].Page1[0].T7[0]":  first_name,
        "topmostSubform[0].Page1[0].T49[0]": tourist.get("passport_number", ""),
        "topmostSubform[0].Page1[0].T50[0]": tourist.get("nationality_iso", "RUS"),
        "topmostSubform[0].Page1[0].T34[0]": "  ",  # Former nationalities — dropdown only accepts country names; blank
        "topmostSubform[0].Page1[0].T37[0]": "NO",  # ID No. issued by government — always NO for Russians
        "topmostSubform[0].Page1[0].#area[4].T14[0]": tourist.get("birth_date", ""),
        "topmostSubform[0].Page1[0].#area[4].T16[0]": tourist.get("place_of_birth", ""),
        "topmostSubform[0].Page1[0].#area[5].#area[6].#area[7].RB1[0]": tourist.get("gender_rb", "0"),
        "topmostSubform[0].Page1[0].#area[8].RB2[0]": tourist.get("marital_status_rb", "0"),
        "topmostSubform[0].Page1[0].#area[1].#area[2].RB3[0]": tourist.get("passport_type_rb", "2"),
        "topmostSubform[0].Page1[0].#area[9].T53[0]": tourist.get("issue_date", ""),
        "topmostSubform[0].Page1[0].#area[0].T59[0]": tourist.get("expiry_date", ""),  # fixed: was T59[0]
        "topmostSubform[0].Page1[0].#area[0].T57[0]": tourist.get("issued_by", ""),       # Issuing Authority = raw MVD code
        "topmostSubform[0].Page1[0].#area[9].T57[1]": _mvd_to_place_of_issue(tourist.get("issued_by", "")),  # Place of Issue = Russia, City
        "topmostSubform[0].Page1[0].T5[0]":  tourist.get("occupation", ""),
        "topmostSubform[0].Page1[0].#area[3].emp_name[0]": tourist.get("employer", ""),
        "topmostSubform[0].Page1[0].emp_adr[0]":           tourist.get("employer_address", ""),
        "topmostSubform[0].Page1[0].#area[3].emp_tel[0]":  tourist.get("employer_phone", ""),
        "topmostSubform[0].Page1[0].T0[1]":  tourist.get("home_address", ""),
        "topmostSubform[0].Page1[0].#area[11].T3[0]": tourist.get("phone", ""),  # applicant phone
        "topmostSubform[0].Page1[0].T62[0]": "NO",                      # Certificate of Eligibility No. — always NO
        "topmostSubform[0].Page1[0].T64[0]": previous_visits_val,       # Dates/duration of previous stays in Japan
        "topmostSubform[0].Page1[0].T66[0]": tourist.get("arrival_date_japan", ""),
        "topmostSubform[0].Page1[0].#area[10].T68[0]": tourist.get("arrival_airport", ""),
        "topmostSubform[0].Page1[0].#area[10].T68[1]": tourist.get("arrival_flight", ""),
        "topmostSubform[0].Page1[0].T68[2]": "tourism",
        "topmostSubform[0].Page1[0].T68[3]": str(stay_days),
        "topmostSubform[0].Page1[0].emp_adr[1]":           anketa.get("first_hotel_address", ""),
        "topmostSubform[0].Page1[0].#area[12].emp_name[1]": anketa.get("first_hotel_name", ""),
        "topmostSubform[0].Page1[0].#area[12].emp_tel[1]":  anketa.get("first_hotel_phone", ""),
        "topmostSubform[0].Page1[0].T3[1]": anketa.get("email", "tour@fujitravel.ru"),
        "topmostSubform[0].Page1[0].T16[1]": maiden_name_val,  # Other names — maiden_name or NO
        # Page 2
        "topmostSubform[0].Page2[0].T150[0]": anketa.get("date_of_application", ""),
        "topmostSubform[0].Page2[0].#area[4].RB5[0]":         anketa.get("criminal_rb", "1"),
        "topmostSubform[0].Page2[0].#area[5].RB5[1]":         anketa.get("criminal_rb", "1"),
        "topmostSubform[0].Page2[0].#area[6].RB5[2]":         anketa.get("criminal_rb", "1"),
        "topmostSubform[0].Page2[0].#area[7].#area[8].RB5[3]": anketa.get("criminal_rb", "1"),
        "topmostSubform[0].Page2[0].#area[9].RB5[4]":         anketa.get("criminal_rb", "1"),
        "topmostSubform[0].Page2[0].RB5[5]":                  anketa.get("criminal_rb", "1"),
        # Guarantor fields — always dash
        "topmostSubform[0].Page2[0].guarantor_adr[0]":           "—",
        "topmostSubform[0].Page2[0].#area[0].guarantor_name[0]": "—",
        "topmostSubform[0].Page2[0].#area[0].guarantor_tel[0]":  "—",
        "topmostSubform[0].Page2[0].#area[1].T14[0]": "—",
        "topmostSubform[0].Page2[0].T25[0]":           "—",
        "topmostSubform[0].Page2[0].T23[0]":           "—",
        "topmostSubform[0].Page2[0].#area[2].T19[0]":  "—",
        "topmostSubform[0].Page2[0].#area[2].T10[0]":  "—",
        "topmostSubform[0].Page2[0].#area[3].T14[1]":  "—",
        "topmostSubform[0].Page2[0].T25[1]":           "—",
        "topmostSubform[0].Page2[0].T28[0]":           "—",
        "topmostSubform[0].Page2[0].T5[0]":  "—",
        "topmostSubform[0].Page2[0].T5[1]":  "—",
        "topmostSubform[0].Page2[0].T5[2]":  "—",
        "topmostSubform[0].Page2[0].T5[3]":  "—",
        "topmostSubform[0].Page2[0].T16[2]": "—",
    }

    fillpdfs.write_fillable_pdf(PDF_TEMPLATE, out_path, fields)

    # Post-pass: overwrite T34 (Former nationalities) with the real value
    # ("USSR" or "NO" — see assembler's ComputeFormerNationality). fillpdf
    # rejects strings outside the dropdown's option list with KeyError;
    # pikepdf's set_value accepts any string because the field is an
    # editable combo box (Ff=393216 = Combo+Edit per the PDF spec).
    # Visible in browsers and Adobe Acrobat. Mac Preview's PDFKit does
    # not render combo values outside the predefined list — that's a
    # known PDFKit limitation, not a bug in this code.
    former = tourist.get("former_nationality_text", "NO")
    try:
        pdf = pikepdf.open(out_path, allow_overwriting_input=True)
        try:
            for f in pdf.acroform.get_fields_with_qualified_name(
                "topmostSubform[0].Page1[0].T34[0]"
            ):
                f.set_value(former)
            pdf.save(out_path)
        finally:
            pdf.close()
    except Exception:
        # Defensive: if the patch fails for any reason, leave fillpdf's
        # safe placeholder in place rather than break document generation.
        # Any failure here is a rendering issue, not a data-loss one —
        # the rest of the anketa is already on disk.
        pass


# ── Main ──────────────────────────────────────────────────────────────────────

MONTH_NAMES = {
    "01": "января", "02": "февраля", "03": "марта", "04": "апреля",
    "05": "мая",    "06": "июня",    "07": "июля",  "08": "августа",
    "09": "сентября","10": "октября","11": "ноября","12": "декабря",
}


def main():
    # Usage: generate.py <group_id> <pass2_json_path> <zip_path> [mode] [subgroup_name]
    # mode: "tourists" (default) | "final"
    # subgroup_name: when set, files go into docs/{subgroup_name}/ and no ZIP is created
    if len(sys.argv) < 4:
        print(f"Usage: {sys.argv[0]} <group_id> <pass2_json_path> <zip_path> [tourists|final] [subgroup_name]")
        sys.exit(1)

    group_id       = sys.argv[1]
    json_path      = sys.argv[2]
    zip_path       = sys.argv[3]
    mode           = sys.argv[4] if len(sys.argv) > 4 else "tourists"
    subgroup_name  = sys.argv[5] if len(sys.argv) > 5 else ""

    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    tourists   = data["tourists"]
    doverenost = data.get("doverenost", [])
    anketa     = data.get("anketa", {})

    base_docs_dir = os.path.join(os.path.dirname(json_path), "docs")
    # If subgroup_name is given, put files in docs/{subgroup_name}/
    if subgroup_name:
        out_dir = os.path.join(base_docs_dir, subgroup_name)
    else:
        out_dir = base_docs_dir
    os.makedirs(out_dir, exist_ok=True)

    generated_files = []

    if mode == "tourists":
        group_label = subgroup_name or "Группа"

        # ── Find contact phone: oldest male, else oldest female ──────────────
        def _birth_key(t):
            bd = t.get("birth_date", "")
            try:
                d, m, y = bd.split(".")
                return (int(y), int(m), int(d))   # sort ascending = oldest first
            except Exception:
                return (9999, 12, 31)

        males   = [t for t in tourists if t.get("gender", "").upper() == "M"]
        females = [t for t in tourists if t.get("gender", "").upper() == "F"]
        if males:
            contact_phone = min(males, key=_birth_key).get("phone", "")
        elif females:
            contact_phone = min(females, key=_birth_key).get("phone", "")
        else:
            contact_phone = tourists[0].get("phone", "") if tourists else ""

        # ── Create output subfolders ─────────────────────────────────────────
        prog_dir = os.path.join(out_dir, "Программы")
        dov_dir  = os.path.join(out_dir, "Доверенности")
        ank_dir  = os.path.join(out_dir, "Анкеты")
        for d in [prog_dir, dov_dir, ank_dir]:
            os.makedirs(d, exist_ok=True)

        # ── ONE programme for the whole group ────────────────────────────────
        prog_path = os.path.join(prog_dir, f"{group_label} прг.docx")
        generate_programme(data, prog_path, contact_phone)
        generated_files.append(prog_path)

        # ── Per-tourist: доверенность + анкета ──────────────────────────────
        for i, tourist in enumerate(tourists):
            name_ru = tourist.get("name_cyr", "")
            parts_ru = name_ru.split()
            surname_ru = parts_ru[0] if parts_ru else f"tourist_{i}"
            first_ru   = " ".join(parts_ru[1:]) if len(parts_ru) > 1 else ""
            file_base_ru = f"{surname_ru} {first_ru}".strip()

            dov_data = doverenost[i] if i < len(doverenost) else {}
            if dov_data:
                # Always name by the tourist's own nominative name (child or adult)
                dov_path = os.path.join(dov_dir, f"{file_base_ru} двр.docx")
                generate_doverenost(data, dov_data, dov_path)
                generated_files.append(dov_path)

            ank_path = os.path.join(ank_dir, f"{file_base_ru} анк.pdf")
            dep_date_str = data.get("departure", {}).get("date", "")
            generate_anketa(tourist, anketa, dov_data, ank_path, dep_date_str)
            generated_files.append(ank_path)

    elif mode == "final":
        # Group-level docs: для Инны в ВЦ + заявка ВЦ
        inna_path = os.path.join(out_dir, "для Инны в ВЦ.docx")
        generate_inna(data, inna_path)
        generated_files.append(inna_path)

        arrival_date = data.get("arrival", {}).get("date", "")
        if arrival_date and len(arrival_date) == 10:
            dd, mm = arrival_date[:2], arrival_date[3:5]
            month_ru = MONTH_NAMES.get(mm, mm)
            vc_filename = f"на {dd} {month_ru} {len(tourists)}.docx"
        else:
            vc_filename = f"заявка ВЦ {len(tourists)}.docx"

        vc_path = os.path.join(out_dir, vc_filename)
        generate_vc_request(data, vc_path)
        generated_files.append(vc_path)

    # When subgroup_name is set, skip ZIP — Go will collect all subgroups and zip them.
    if subgroup_name:
        print(f"Generated {len(generated_files)} files for subgroup '{subgroup_name}' (no zip)")
        return

    # Zip everything (no subgroups — single folder mode).
    os.makedirs(os.path.dirname(zip_path) or ".", exist_ok=True)
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for fp in generated_files:
            if os.path.exists(fp):
                zf.write(fp, os.path.basename(fp))

    print(f"Generated {len(generated_files)} files → {zip_path}")


if __name__ == "__main__":
    main()
